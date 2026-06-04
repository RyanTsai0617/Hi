"""AVGO（Broadcom）財報分析 — Word（.docx）底稿產生器。

對應「做財報流程」的最終交付物：Word 底稿（非 PPT）。
沿用 repo 視覺色調（teal #1B4F5A、yellow accent、微軟正黑體），
圖表重用 avgo_financial_report_pptx 的 matplotlib 函式。

章節（標準法人財報分析結構，可再對齊 UWS 12 章節模板）：
  一、摘要   二、公司概況   三、本季財報重點（Q1 FY2026）
  四、營收分析   五、AI 業務分析   六、業務結構（Segment）
  七、獲利能力與現金流   八、財務體質與股東回報
  九、展望與財測   十、投資觀點與風險
  十一、對台灣供應鏈意義   十二、資料來源與聲明

跑法： python3 avgo_financial_report_docx.py  →  輸出 avgo_financial_report.docx
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 重用 pptx 模組的圖表與資料
from avgo_financial_report_pptx import (
    make_revenue_chart, make_ai_revenue_chart, make_segment_donut, make_fcf_chart,
    QUARTERS,
)

ZH_FONT = "微軟正黑體"
EN_FONT = "Calibri"
TEAL = RGBColor(0x1B, 0x4F, 0x5A)
TEAL_LIGHT = RGBColor(0x2B, 0x7A, 0x78)
GREY = RGBColor(0x55, 0x55, 0x55)
YELLOW = "E8C547"
TEAL_HEX = "1B4F5A"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------- helpers ----------------

def _set_run(run, *, size=11, bold=False, color=None, font_zh=ZH_FONT, font_en=EN_FONT):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_en
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_zh)


def add_para(doc, text="", *, size=11, bold=False, color=None, align=None,
             space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.25
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, num, text, *, size=15):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{num}　{text}")
    _set_run(run, size=size, bold=True, color=TEAL)
    # teal 底線（段落下框線）
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), TEAL_HEX)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    _set_run(run, size=12, bold=True, color=TEAL_LIGHT)
    return p


def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(it)
        _set_run(run, size=size)


def add_chart(doc, buf, caption, *, width_in=6.2):
    doc.add_picture(buf, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    _set_run(run, size=9, color=GREY)


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, *, header_fill=TEAL_HEX, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        _shade(cell, header_fill)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        _set_run(run, size=10, bold=True, color=WHITE)
    for r, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cell = cells[j]
            if r % 2 == 1:
                _shade(cell, "F2F4F5")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(val))
            color = None
            sval = str(val)
            if sval.startswith("+"):
                color = RGBColor(0x15, 0x80, 0x3D)
            elif sval.startswith("-") and len(sval) > 1 and sval[1].isdigit():
                color = RGBColor(0xB9, 0x1C, 0x1C)
            _set_run(run, size=10, color=color)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)


# ---------------- build ----------------

def build(out_path="avgo_financial_report.docx"):
    doc = Document()
    _set_base_style(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ---- 標題區 ----
    add_para(doc, "Broadcom（AVGO）財報分析", size=24, bold=True, color=TEAL,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "FY2025 全年 ‧ Q1 FY2026 最新季報（截至 2026/02/01）", size=13,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "NASDAQ: AVGO　|　半導體 + 基礎架構軟體　|　AI 客製化加速器（XPU）與 AI 網通領導者",
             size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "國泰證券法人債券業務部 ‧ 法人業務二處　|　報告日期：2026/06/04",
             size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # ---- 一、摘要 ----
    add_heading(doc, "一、", "摘要（Executive Summary）")
    add_bullets(doc, [
        "AI 為核心成長引擎：FY2025 AI 營收達 $20B，Q1 FY2026 單季 AI 半導體營收 $8.4B、YoY +106%，"
        "佔總營收比重由一年前約 27% 拉升至約 44%。",
        "營運規模續創高：FY2025 全年合併營收 $62.9B（+22% YoY）；Q1 FY2026 營收 $19.31B（+29% YoY），全面優於財測。",
        "高獲利、高現金流：Q1 FY2026 營業利益率 66.4%、Adjusted EBITDA 佔營收 68%、自由現金流 $8.0B（佔營收 41%）。",
        "能見度明確：Q2 FY2026 財測約 $22.0B（+47% YoY）；管理層對 2027 年 AI 晶片營收 line of sight 上看 $100B+，"
        "AI backlog 約 $73B、交期能見度約 18 個月。",
        "雙引擎結構：半導體（AI 高成長）+ 基礎架構軟體（VMware，高毛利訂閱現金流）並進。",
    ])

    # ---- 二、公司概況 ----
    add_heading(doc, "二、", "公司概況")
    add_para(doc, "Broadcom Inc.（NASDAQ: AVGO）為全球領先的半導體與基礎架構軟體供應商，"
                  "業務分為兩大 segment：")
    add_bullets(doc, [
        "半導體解決方案（Semiconductor Solutions）：含 AI 客製化加速器（XPU/ASIC）、AI 網通"
        "（Tomahawk / Jericho 交換器晶片、光通訊）、無線、寬頻、伺服器儲存等。",
        "基礎架構軟體（Infrastructure Software）：以併購 VMware 為核心，提供虛擬化、雲端與企業軟體訂閱。",
    ])
    add_para(doc, "公司營運模式以高市佔的關鍵元件 + 高黏著度的軟體訂閱組合，創造強勁且可預期的現金流，"
                  "並以高比例配息與去槓桿回饋股東。")

    # ---- 三、本季財報重點 ----
    add_heading(doc, "三、", "本季財報重點（Q1 FY2026，2026/03/04 公布）")
    add_table(doc,
              ["指標", "本季數值", "YoY / 備註"],
              [
                  ["合併營收", "$19.31B", "+29%"],
                  ["半導體營收（紀錄）", "$12.5B", "+52%"],
                  ["AI 半導體營收", "$8.4B", "+106%"],
                  ["基礎架構軟體營收", "約 $6.8B", "約 +1%"],
                  ["GAAP 淨利", "$7.35B", "GAAP EPS $1.50"],
                  ["Non-GAAP 淨利", "$10.19B", "Non-GAAP EPS $2.05"],
                  ["營業利益率", "66.4%", "+50 bps"],
                  ["Adjusted EBITDA", "$13.1B", "佔營收 68%"],
                  ["營運現金流", "$8.26B", "—"],
                  ["自由現金流", "$8.01B", "佔營收 41%"],
                  ["季配息", "$0.65 / 股", "—"],
              ],
              widths=[2.4, 1.8, 2.2])
    add_para(doc, "重點：AI 加速器（XPU）＋ AI 網通雙引擎帶動半導體單季創高；軟體（VMware）提供穩定現金流。",
             size=10, color=TEAL, bold=True)

    # ---- 四、營收分析 ----
    add_heading(doc, "四、", "營收分析")
    add_chart(doc, make_revenue_chart(width_in=9.5, height_in=3.7),
              "圖 1　近 5 季合併營收（長條）與 YoY 成長率（折線）")
    rows = [[q[0], f"${q[1]/1000:.1f}B", f"+{q[3]}%", f"${q[2]/1000:.1f}B"] for q in QUARTERS]
    add_table(doc, ["季別", "合併營收", "YoY", "其中 AI 營收"], rows,
              widths=[1.6, 1.8, 1.4, 1.8])
    add_para(doc, "FY2025 全年營收 $62.9B（+22% YoY，FY2024 為 $51.6B），季營收連續創高、YoY 由 +20% 區間"
                  "加速至 Q1 FY2026 的 +29%，動能主要來自 AI 半導體。")

    # ---- 五、AI 業務分析 ----
    add_heading(doc, "五、", "AI 業務分析")
    add_chart(doc, make_ai_revenue_chart(width_in=9.5, height_in=3.7),
              "圖 2　AI 半導體營收（季）與佔總營收比重")
    add_bullets(doc, [
        "成長軌跡：AI 半導體營收由 Q1'25 約 $4.1B 成長至 Q1'26 的 $8.4B（YoY +106%）。",
        "雙產品線：custom 加速器（XPU/ASIC）＋ AI 網通（Ethernet switch / 光通訊）。",
        "中長期：管理層對 2027 年 AI 晶片營收 line of sight 上看 $100B+；AI backlog 約 $73B、約 18 個月交期。",
        "Q2 FY2026 AI 半導體營收預估約 $10.7B，續創高。",
    ])

    # ---- 六、業務結構 ----
    add_heading(doc, "六、", "業務結構（Segment，FY2025）")
    add_chart(doc, make_segment_donut(width_in=4.6, height_in=3.8),
              "圖 3　FY2025 兩大 segment 營收結構（約略值）", width_in=4.3)
    add_para(doc, "FY2025 半導體解決方案約 $37B（約 59%），基礎架構軟體約 $26B（約 41%）。"
                  "半導體提供 AI 高成長，軟體（VMware 整併）提供高毛利、可預期的訂閱現金流，"
                  "兩者形成成長 + 現金流的互補結構。")

    # ---- 七、獲利能力與現金流 ----
    add_heading(doc, "七、", "獲利能力與現金流")
    add_chart(doc, make_fcf_chart(width_in=5.2, height_in=3.7),
              "圖 4　FY 自由現金流 vs GAAP 淨利", width_in=5.0)
    add_bullets(doc, [
        "獲利率：Q1 FY2026 營業利益率 66.4%（+50 bps YoY）、Adjusted EBITDA 佔營收 68%。",
        "現金流：FY2025 自由現金流 $26.9B；Q1 FY2026 單季 FCF $8.0B（佔營收 41%），資本支出僅約 $250M（輕資產）。",
        "獲利躍升：GAAP 淨利由 FY2024 的 $5.9B（受 VMware 併購相關費用壓抑）回升至 FY2025 的 $23.1B。",
    ])

    # ---- 八、財務體質與股東回報 ----
    add_heading(doc, "八、", "財務體質與股東回報")
    add_bullets(doc, [
        "強勁 FCF 支撐穩定配息（Q1 FY2026 季配 $0.65/股）與持續去槓桿（償還 VMware 併購相關負債）。",
        "輕資本支出模式：以 fabless 半導體 + 軟體訂閱為主，capex 佔營收比重低，現金轉換效率高。",
        "高毛利結構：軟體訂閱與高市佔的關鍵元件，支撐長期毛利率與營益率維持高檔。",
    ])

    # ---- 九、展望與財測 ----
    add_heading(doc, "九、", "展望與財測")
    add_table(doc, ["項目", "內容"],
              [
                  ["Q2 FY2026 營收財測", "約 $22.0B（+47% YoY）"],
                  ["Q2 AI 半導體營收", "預估約 $10.7B"],
                  ["2027 AI 晶片營收", "line of sight 上看 $100B+"],
                  ["AI backlog", "約 $73B、約 18 個月交期"],
                  ["獲利率展望", "營益率維持 66%+ 高檔"],
              ],
              widths=[2.6, 4.0])

    # ---- 十、投資觀點與風險 ----
    add_heading(doc, "十、", "投資觀點與風險")
    add_subheading(doc, "投資觀點")
    add_bullets(doc, [
        "AI 結構性成長 + 軟體現金流的稀有組合，AI 營收能見度高（backlog + 2027 目標）。",
        "可作為 AI 半導體 ETF（SMH / SOXX / AIQ）權重與題材的核心觀察錨。",
    ])
    add_subheading(doc, "主要風險")
    add_bullets(doc, [
        "客製化 ASIC 客戶集中度高，單一超大型客戶資本支出變動影響大。",
        "基礎架構軟體成長趨緩（Q1 約 +1%），未來成長更倚賴 AI 半導體。",
        "評價偏高，對 AI 資本支出循環、利率與總經風險敏感。",
        "地緣政治與出口管制對先進製程供應鏈的潛在影響。",
    ])

    # ---- 十一、對台灣供應鏈意義 ----
    add_heading(doc, "十一、", "對台灣供應鏈意義")
    add_bullets(doc, [
        "AVGO 為台積電（TSMC）先進製程（CoWoS / N3 / N2）核心大客戶之一，XPU 放量直接拉動先進封裝需求。",
        "AI 加速器放量 → 台系 ABF 載板、CCL、測試、散熱供應鏈受惠。",
        "AI 網通（Tomahawk / Jericho）利多台系交換器、光通訊（CPO / 光模組）族群。",
        "可作為觀察台股 AI 供應鏈景氣與 SMH/SOXX 連動的領先指標之一。",
    ])

    # ---- 十二、資料來源與聲明 ----
    add_heading(doc, "十二、", "資料來源與聲明")
    add_bullets(doc, [
        "Broadcom Inc.「Q1 FY2026 Financial Results」官方新聞稿（2026/03/04）。",
        "Broadcom Inc.「Q4 & FY2025 Financial Results」官方新聞稿（2025/12/11）。",
        "SEC Form 8-K 財報附件（CIK 0001730168，FY2025 / FY2026 各季）。",
        "全年數字部分由各季 SEC 8-K 加總推算；segment 為約略值，以官方 10-K 為準。",
    ], size=10)
    add_para(doc, "本報告僅為財報數據彙整與分析，供內部研究與法人業務參考，不構成任何投資建議或要約；"
                  "投資人應自行評估風險。金額單位為美元（US$），會計期間以 Broadcom 財報年度"
                  "（FY 結束於 11 月初）為準。", size=10, color=GREY)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[OK] AVGO financial report (Word) saved: {out_path}  ({len(doc.paragraphs)} paragraphs)")
    return out_path


if __name__ == "__main__":
    build()
